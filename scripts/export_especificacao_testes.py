#!/usr/bin/env python3
"""Gera o arquivo EspecificacaoTestes.docx com base no conteúdo do componente React."""

from pathlib import Path

from docx import Document


def add_subtitle(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered_list(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def main() -> None:
    project_root = Path(__file__).resolve().parents[1]
    output_dir = project_root / "docs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "EspecificacaoTestes.docx"

    document = Document()

    document.add_heading("Especificação de Testes - Sistema SaúdeJá", level=0)
    add_subtitle(
        document,
        "Documento de Especificação de Testes do Módulo de Agendamento de Consultas",
    )

    document.add_heading("1. Informações Gerais", level=1)
    info_table = document.add_table(rows=4, cols=2)
    info_table.style = "Light List Accent 1"
    info_table.cell(0, 0).text = "Sistema"
    info_table.cell(0, 1).text = "SaúdeJá - Plataforma de Agendamento de Consultas"
    info_table.cell(1, 0).text = "Módulo Testado"
    info_table.cell(1, 1).text = (
        "Módulo de Agendamento, Alteração e Cancelamento de Consultas"
    )
    info_table.cell(2, 0).text = "Versão"
    info_table.cell(2, 1).text = "1.0.0 - Protótipo"
    info_table.cell(3, 0).text = "Data"
    info_table.cell(3, 1).text = "Novembro de 2025"

    objective_paragraph = document.add_paragraph()
    objective_paragraph.add_run("Objetivo dos Testes: ").bold = True
    objective_paragraph.add_run(
        "Validar o funcionamento correto do módulo de agendamento de consultas, "
        "garantindo que os usuários possam agendar, alterar e cancelar consultas de "
        "forma eficiente e sem erros, com notificações adequadas e validações de dados."
    )

    document.add_heading("2. Tipos de Testes", level=1)
    document.add_paragraph("Categorias de testes que serão aplicados ao sistema.")
    document.add_heading("Testes Funcionais", level=2)
    document.add_paragraph(
        "Verificação de funcionalidades principais: agendamento, alteração, cancelamento, "
        "validações de formulário e persistência de dados."
    )
    document.add_heading("Testes de Usabilidade", level=2)
    document.add_paragraph(
        "Avaliação da interface, navegação intuitiva, clareza das mensagens e facilidade de uso "
        "para diferentes perfis de usuários."
    )
    document.add_heading("Testes de Validação", level=2)
    document.add_paragraph(
        "Verificação de regras de negócio: campos obrigatórios, formatos de dados, datas válidas "
        "e prevenção de agendamentos duplicados."
    )
    document.add_heading("Testes de Integração", level=2)
    document.add_paragraph(
        "Verificação de integração entre componentes: formulário ↔ lista de consultas, "
        "persistência de dados e sistema de notificações."
    )

    document.add_heading("3. Módulos a Serem Testados", level=1)
    document.add_paragraph("Componentes e funcionalidades específicas do sistema.")

    document.add_heading("Módulo de Agendamento de Consultas", level=2)
    document.add_paragraph("Componentes:")
    add_bullet_list(
        document,
        [
            "Formulário de agendamento",
            "Seleção de unidade de saúde",
            "Seleção de especialidade e profissional",
            "Calendário de datas disponíveis",
            "Seleção de horários",
            "Sistema de notificações",
            "Persistência de dados (localStorage)",
        ],
    )

    document.add_heading("Módulo de Alteração de Consultas", level=2)
    document.add_paragraph("Componentes:")
    add_bullet_list(
        document,
        [
            "Listagem de consultas agendadas",
            "Modal de reagendamento",
            "Seleção de nova data",
            "Seleção de novo horário",
            "Atualização de dados persistidos",
            "Confirmação de alteração",
            "Notificação de reagendamento",
        ],
    )

    document.add_heading("Módulo de Cancelamento de Consultas", level=2)
    document.add_paragraph("Componentes:")
    add_bullet_list(
        document,
        [
            "Botão de cancelamento",
            "Dialog de confirmação",
            "Atualização de status da consulta",
            "Movimentação para histórico de canceladas",
            "Notificação de cancelamento",
            "Atualização de dados persistidos",
        ],
    )

    document.add_heading("4. Casos de Teste Detalhados", level=1)

    def add_test_case(
        code: str,
        title: str,
        category: str,
        objective: str,
        preconditions: str | None,
        steps: list[str],
        expected_result: str,
    ) -> None:
        document.add_heading(f"{code}: {title}", level=2)
        badge_paragraph = document.add_paragraph()
        badge_paragraph.add_run(f"Categoria: {category}")
        objective_paragraph = document.add_paragraph()
        objective_paragraph.add_run("Objetivo: ").bold = True
        objective_paragraph.add_run(objective)
        if preconditions:
            preconditions_paragraph = document.add_paragraph()
            preconditions_paragraph.add_run("Pré-condições: ").bold = True
            preconditions_paragraph.add_run(preconditions)
        document.add_paragraph("Passos:")
        add_numbered_list(document, steps)
        expected_paragraph = document.add_paragraph()
        expected_paragraph.add_run("Resultado Esperado: ").bold = True
        expected_paragraph.add_run(expected_result)

    add_test_case(
        code="CT-001",
        title="Agendamento de Consulta com Sucesso",
        category="Funcional",
        objective=(
            "Verificar se o sistema permite agendar uma consulta preenchendo todos os campos corretamente."
        ),
        preconditions="Sistema carregado e formulário de agendamento visível.",
        steps=[
            "Preencher nome do paciente.",
            "Selecionar unidade de saúde.",
            "Selecionar especialidade.",
            "Selecionar profissional.",
            "Selecionar data futura.",
            "Selecionar horário disponível.",
            'Clicar em "Confirmar Agendamento".',
        ],
        expected_result=(
            "Mensagem de sucesso exibida, consulta salva, notificação enviada e formulário limpo."
        ),
    )

    add_test_case(
        code="CT-002",
        title="Tentativa de Agendamento com Campos Vazios",
        category="Validação",
        objective=(
            "Verificar se o sistema valida campos obrigatórios e impede agendamento incompleto."
        ),
        preconditions=None,
        steps=[
            "Deixar um ou mais campos em branco.",
            'Clicar em "Confirmar Agendamento".',
        ],
        expected_result=(
            "Mensagem de erro exibida, agendamento não realizado e campos permanecem preenchidos."
        ),
    )

    add_test_case(
        code="CT-003",
        title="Reagendamento de Consulta",
        category="Funcional",
        objective=(
            "Verificar se o sistema permite alterar data e horário de uma consulta agendada."
        ),
        preconditions="Pelo menos uma consulta agendada no sistema.",
        steps=[
            'Navegar para "Minhas Consultas".',
            "Clicar no botão de editar em uma consulta.",
            "Selecionar nova data.",
            "Selecionar novo horário.",
            "Confirmar reagendamento.",
        ],
        expected_result=(
            "Consulta atualizada com nova data/horário, mensagem de sucesso e notificação enviada."
        ),
    )

    add_test_case(
        code="CT-004",
        title="Cancelamento de Consulta",
        category="Funcional",
        objective="Verificar se o sistema permite cancelar uma consulta agendada.",
        preconditions="Pelo menos uma consulta agendada no sistema.",
        steps=[
            'Navegar para "Minhas Consultas".',
            "Clicar no botão de cancelar em uma consulta.",
            "Confirmar cancelamento no dialog.",
        ],
        expected_result=(
            'Status alterado para "cancelada", consulta movida para histórico e notificação enviada.'
        ),
    )

    add_test_case(
        code="CT-005",
        title="Persistência de Dados",
        category="Integração",
        objective="Verificar se as consultas são persistidas e recuperadas corretamente.",
        preconditions=None,
        steps=[
            "Agendar uma nova consulta.",
            "Recarregar a página.",
            'Navegar para "Minhas Consultas".',
        ],
        expected_result="Consulta continua visível na lista com todos os dados corretos.",
    )

    add_test_case(
        code="CT-006",
        title="Validação de Data Passada",
        category="Validação",
        objective="Verificar se o sistema impede seleção de datas no passado.",
        preconditions=None,
        steps=[
            "Abrir calendário de seleção de data.",
            "Tentar selecionar uma data anterior à data atual.",
        ],
        expected_result="Datas passadas desabilitadas ou não selecionáveis no calendário.",
    )

    add_test_case(
        code="CT-007",
        title="Sistema de Notificações",
        category="Funcional",
        objective="Verificar se notificações são exibidas nas ações principais.",
        preconditions=None,
        steps=[
            "Agendar uma consulta.",
            "Reagendar uma consulta.",
            "Cancelar uma consulta.",
        ],
        expected_result=(
            "Toast de confirmação exibido em cada ação e notificação de e-mail simulada."
        ),
    )

    add_test_case(
        code="CT-008",
        title="Seleção Condicional de Profissionais",
        category="Usabilidade",
        objective=(
            "Verificar se a lista de profissionais é filtrada pela especialidade selecionada."
        ),
        preconditions=None,
        steps=[
            "Selecionar uma especialidade (ex.: Cardiologia).",
            "Abrir o select de profissionais.",
            "Verificar profissionais listados.",
            "Mudar para outra especialidade.",
            "Verificar que a lista de profissionais mudou.",
        ],
        expected_result=(
            "Apenas profissionais da especialidade selecionada são exibidos."
        ),
    )

    document.add_heading("5. Critérios de Aceitação", level=1)
    criteria = [
        (
            "Taxa de Sucesso:",
            "100% dos casos de teste principais (CT-001 a CT-008) devem passar.",
        ),
        (
            "Validações:",
            "Todos os campos obrigatórios devem ser validados corretamente.",
        ),
        (
            "Notificações:",
            "Sistema deve exibir feedback visual para todas as ações do usuário.",
        ),
        (
            "Persistência:",
            "Dados devem ser mantidos entre recarregamentos da página.",
        ),
        (
            "Usabilidade:",
            "Interface deve ser intuitiva e responsiva.",
        ),
        (
            "Responsividade:",
            "Sistema deve funcionar corretamente em diferentes tamanhos de tela.",
        ),
    ]
    for title, description in criteria:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{title} ").bold = True
        paragraph.add_run(description)

    document.add_heading("6. Matriz de Rastreabilidade", level=1)
    rows = [
        ("RF-01: Agendar consulta", "CT-001, CT-002, CT-005, CT-006", "Alta"),
        ("RF-02: Alterar consulta", "CT-003, CT-005", "Alta"),
        ("RF-03: Cancelar consulta", "CT-004, CT-005", "Alta"),
        ("RF-04: Enviar notificações", "CT-007", "Média"),
        ("RF-05: Filtrar profissionais por especialidade", "CT-008", "Média"),
        ("RF-06: Validar campos obrigatórios", "CT-002", "Alta"),
        ("RF-07: Persistir dados", "CT-005", "Alta"),
        ("RF-08: Impedir datas passadas", "CT-006", "Média"),
    ]
    table = document.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Requisito Funcional", "Casos de Teste", "Prioridade"]
    for col_index, header in enumerate(headers):
        table.cell(0, col_index).text = header
    for row_index, (requirement, test_cases, priority) in enumerate(rows, start=1):
        table.cell(row_index, 0).text = requirement
        table.cell(row_index, 1).text = test_cases
        table.cell(row_index, 2).text = priority

    document.add_heading("7. Ambiente de Testes", level=1)
    section_data = [
        (
            "Tecnologias:",
            [
                "React + TypeScript",
                "Tailwind CSS",
                "Shadcn/UI Components",
                "LocalStorage para persistência",
            ],
        ),
        (
            "Navegadores Suportados:",
            [
                "Google Chrome (última versão)",
                "Mozilla Firefox (última versão)",
                "Safari (última versão)",
                "Microsoft Edge (última versão)",
            ],
        ),
        (
            "Dispositivos:",
            [
                "Desktop (1920x1080 e superiores)",
                "Tablet (768x1024)",
                "Mobile (375x667 e superiores)",
            ],
        ),
        (
            "Ferramentas de Teste:",
            [
                "DevTools do navegador",
                "Testes manuais",
                "Simulação de diferentes resoluções",
            ],
        ),
    ]
    for title, items in section_data:
        paragraph = document.add_paragraph()
        paragraph.add_run(title).bold = True
        add_bullet_list(document, items)

    document.add_heading("8. Considerações Finais", level=1)
    final_notes = [
        (
            "Limitações do Protótipo:",
            "Este protótipo utiliza armazenamento local (localStorage) para simular persistência de dados. "
            "Em um ambiente de produção, seria necessária integração com backend e banco de dados.",
        ),
        (
            "Próximas Etapas:",
            "Implementação de testes automatizados (Jest, Testing Library), testes de carga, integração com API real e testes de segurança.",
        ),
        (
            "Observações:",
            "As notificações de e-mail e SMS são simuladas neste protótipo. Em produção, seria necessária integração com serviços de envio de mensagens.",
        ),
    ]
    for title, content in final_notes:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{title} ").bold = True
        paragraph.add_run(content)

    document.save(output_path)
    print(f"Documento gerado em: {output_path}")


if __name__ == "__main__":
    main()

